import base64
import io
import json
import os
from pathlib import Path

import pandas as pd
from docx import Document
from PIL import Image
from pypdf import PdfReader

from .file_engine import detect_file_type


def _clean(text: str) -> str:
    return "\n".join(line.strip() for line in (text or "").splitlines() if line.strip())


def parse_pdf(path: str | Path) -> dict:
    reader = PdfReader(str(path))
    pages = []
    for idx, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        pages.append({"page": idx + 1, "text": _clean(text)})
    combined = "\n\n".join(f"[PAGE {p['page']}]\n{p['text']}" for p in pages if p["text"])
    return {
        "text": combined,
        "meta": {"pages": len(reader.pages), "text_pages": sum(bool(p["text"]) for p in pages)},
        "needs_vision": len(combined.strip()) < 80,
    }


def parse_word(path: str | Path) -> dict:
    if Path(path).suffix.lower() == ".doc":
        return {"text": "", "meta": {"warning": "Legacy .doc requires conversion to .docx for local parsing."}, "needs_vision": False}
    doc = Document(str(path))
    chunks = []
    for p in doc.paragraphs:
        if p.text.strip():
            chunks.append(p.text.strip())
    for t_idx, table in enumerate(doc.tables, start=1):
        chunks.append(f"[TABLE {t_idx}]")
        for row in table.rows:
            chunks.append(" | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells))
    return {"text": _clean("\n".join(chunks)), "meta": {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}, "needs_vision": False}


def parse_excel(path: str | Path) -> dict:
    ext = Path(path).suffix.lower()
    chunks = []
    sheets_meta = []
    if ext == ".csv":
        df = pd.read_csv(path)
        frames = {"CSV": df}
    else:
        frames = pd.read_excel(path, sheet_name=None)
    for sheet, df in frames.items():
        df = df.dropna(how="all").dropna(axis=1, how="all")
        sheets_meta.append({"sheet": str(sheet), "rows": int(df.shape[0]), "columns": int(df.shape[1])})
        chunks.append(f"[SHEET: {sheet}]\n{df.head(200).to_csv(index=False)}")
    return {"text": "\n\n".join(chunks), "meta": {"sheets": sheets_meta}, "needs_vision": False}


def parse_image_local(path: str | Path) -> dict:
    with Image.open(path) as im:
        meta = {"width": im.width, "height": im.height, "format": im.format}
    # We intentionally do not run local OCR here. Vision extraction is performed by the AI layer when configured.
    return {"text": "", "meta": meta, "needs_vision": True}


def parse_document(path: str | Path) -> dict:
    ftype = detect_file_type(path)
    if ftype == "pdf":
        result = parse_pdf(path)
    elif ftype == "word":
        result = parse_word(path)
    elif ftype == "excel":
        result = parse_excel(path)
    elif ftype == "image":
        result = parse_image_local(path)
    else:
        return {"file_type": ftype, "text": "", "meta": {}, "needs_vision": False, "status": "unsupported_local_parser"}
    return {"file_type": ftype, **result, "status": "parsed"}


def image_to_data_url(path: str | Path) -> str:
    suffix = Path(path).suffix.lower().lstrip(".") or "png"
    mime = "jpeg" if suffix in {"jpg", "jpeg"} else suffix
    data = base64.b64encode(Path(path).read_bytes()).decode("ascii")
    return f"data:image/{mime};base64,{data}"


def pdf_pages_as_data_urls(path: str | Path, max_pages: int = 8) -> list[str]:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return []
    doc = fitz.open(str(path))
    out = []
    for idx in range(min(len(doc), max_pages)):
        pix = doc[idx].get_pixmap(matrix=fitz.Matrix(1.4, 1.4), alpha=False)
        data = base64.b64encode(pix.tobytes("png")).decode("ascii")
        out.append(f"data:image/png;base64,{data}")
    doc.close()
    return out
