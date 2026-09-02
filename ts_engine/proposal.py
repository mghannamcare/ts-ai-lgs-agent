from pathlib import Path
from docx import Document
from docx.shared import Pt

def _add_list(doc,title,values):
    doc.add_heading(title,level=2)
    if not values: doc.add_paragraph('Not identified / requires review.'); return
    for v in values:
        text=' — '.join(str(x) for x in v.values() if x not in (None,'',[],{})) if isinstance(v,dict) else str(v)
        doc.add_paragraph(text,style='List Bullet')

def _table(doc,title,rows):
    doc.add_heading(title,level=2)
    if not rows: doc.add_paragraph('No data available.'); return
    cols=list(rows[0].keys()); t=doc.add_table(rows=1, cols=len(cols)); t.style='Table Grid'
    for i,c in enumerate(cols): t.rows[0].cells[i].text=str(c)
    for r in rows[:250]:
        cells=t.add_row().cells
        for i,c in enumerate(cols): cells[i].text=str(r.get(c,''))[:1000]

def build_technical_proposal(analysis, output_path, compliance=None):
    doc=Document(); doc.styles['Normal'].font.name='Arial'; doc.styles['Normal'].font.size=Pt(10)
    doc.add_heading('TS Technical Proposal – AI Draft',0)
    doc.add_paragraph(f"Project: {analysis.get('project','')}"); doc.add_paragraph(f"Client: {analysis.get('client','')}")
    doc.add_paragraph('AI-assisted draft. TS technical and commercial approval is required before issue.')
    doc.add_heading('Executive Summary',1); doc.add_paragraph(str(analysis.get('executive_summary','')))
    _add_list(doc,'Proposed Scope of Work',analysis.get('scope_of_work',[])); _table(doc,'BOQ / Quantity Review',analysis.get('boq_summary',[]) or [])
    _add_list(doc,'Technical Requirements',analysis.get('technical_requirements',[])); _add_list(doc,'Proposed TS Solutions',analysis.get('ts_solutions',[]))
    if compliance: _table(doc,'Compliance Matrix',compliance)
    _add_list(doc,'Clarifications / Information Required',analysis.get('clarifications',[])); _add_list(doc,'Risks and Qualifications',analysis.get('risks',[]))
    _table(doc,'Evidence / Source Traceability',analysis.get('evidence',[]) or [])
    d=analysis.get('bid_decision',{}) or {}; doc.add_heading('Bid Assessment',2); doc.add_paragraph(f"Recommendation: {d.get('recommendation','')}"); doc.add_paragraph(f"Score: {d.get('score','')}"); doc.add_paragraph(f"Rationale: {d.get('rationale','')}")
    doc.save(str(output_path)); return str(output_path)
